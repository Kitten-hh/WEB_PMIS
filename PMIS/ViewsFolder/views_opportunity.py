from BaseProject.views_folder.base_crud.SimpleView import SimpleView
from .. forms import Opportunity_main_from
from BaseProject.decorators.customer_dec import URL
from DataBase_MPMS.models import Tecmb, Tecmc, Tecma,VTecmb,VTecma
from django.forms.models import model_to_dict
from django.shortcuts import render
from django.http import HttpResponse, JsonResponse
from BaseApp.library.cust_views.DatatablesServerSideView import DatatablesServerSideView
from BaseApp.library.cust_views.SWUpdateView import SWUpdateView
from BaseApp.library.cust_views.SWCreateView import SWCreateView
from django.db.models import Count,Q,Case,Sum,When,IntegerField,Max,F
from django.db.models.fields import TextField
from DataBase_MPMS import models
import re
import datetime
from django.db.models.signals import pre_save
from django.dispatch import receiver
import base64
from sys import argv
from base64 import b64encode
from json import dumps
from django.contrib.auth.decorators import login_required
import json
from BaseApp.library.tools import ModelTools
from PMIS.Services.ElasticService import ElasticService
from docx import Document
from docx.shared import Pt
# from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from django.core.cache import cache
from django.conf import settings
from BaseApp.library.tools import AsyncioTools
import threading

#將概念轉成Word並下載
def downloadTecnicalWord(request):
    id = request.GET.get("id")
    qs = Tecmb.objects.filter(inc_id = id).first()
    mb004 = qs.mb004 if qs.mb004 is not None else "None" #主題
    # mb005 = qs.mb005 if qs.mb005 is not None else "None" #聯繫人
    # mb006 = qs.mb006.format('%Y-%m-%d') if qs.mb006 is not None else "None" #日期
    # if mb006 != "None":
    #     mb006 = f'{qs.mb006[:4]}-{qs.mb006[4:6]}-{qs.mb006[6:8]}' #將字符串由yyyymmdd轉為yyyy-mm-dd
    mb007 = qs.mb007 if qs.mb007 is not None else "None" #概念
    mb023 = qs.mb023 if qs.mb023 is not None else "None" #文檔編號

    #創建一個新的Word文檔
    doc = Document()
    
    #添加標題
    title = doc.add_heading(qs.mb004, level=1) #設置一級標題
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER #將標題設為居中

    # 添加一个节（section），这是一个页面设置的容器
    section = doc.sections[0]

    # 添加页眉
    header = section.header

    # 在页眉中添加文本
    p1  = header.add_paragraph() #添加一個段落
    run = p1 .add_run('Technical ID: {}'.format(mb023)) #文檔編號
    font = run.font
    font.bold = True #設置字體加粗
    font.size = Pt(12)  # 设置字体大小
    p1.alignment = WD_ALIGN_PARAGRAPH.LEFT  # 设置对齐方式为左对齐

    # paragraph = header.add_paragraph() #添加一個段落
    # run = paragraph.add_run('Contact: {}'.format(mb005)) #聯繫人
    # font = run.font
    # font.bold = True #設置字體加粗
    # font.size = Pt(12)  # 设置字体大小
    # paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT  # 设置对齐方式为左对齐

    # # 添加页眉
    # paragraph = header.add_paragraph() #添加一個段落
    # run = paragraph.add_run('Date: {}'.format(mb006)) #日期 
    # font = run.font
    # font.bold = True #設置字體加粗
    # font.size = Pt(12)
    # paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT  # 设置对齐方式为右对齐

    # #添加標題
    # title = doc.add_heading(qs.mb004,0) #名稱
    # # 获取标题段落对象
    # title_paragraph = title.paragraphs[0]

    # # 设置标题段落的对齐方式为居中
    # title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p2 = doc.add_paragraph() #添加一個段落
    run = p2.add_run("一: Concept/Theory") #設置段落內容,標題為概念
    font = run.font
    font.bold = True #設置字體加粗
    font.size = Pt(16) #設置字體大小為12磅

    p3 = doc.add_paragraph(mb007) #添加一個段落,內容為概念信息
    # 获取段落的格式设置
    paragraph_format = p3.paragraph_format
    # 设置左缩进
    paragraph_format.left_indent = Pt(17)
    # 设置右缩进为0.5英寸
    # paragraph_format.right_indent = Pt(36)  # 0.5英寸 = 36磅

    # 保存文档到内存中
    response = HttpResponse(content_type='application/msword')
    response['Content-Disposition'] = 'attachment; filename={}.docx'.format(mb023)#文檔名稱
    doc.save(response)

    return response
 
## Json數據轉換的處理類
class DataEncoder(json.JSONEncoder):
    def default(self, obj):
        if isinstance(obj, bytes):
            return str(obj,encoding='utf-8')
        return json.JSONEncoder.default(self,obj)


class OpportunityView(SimpleView):
    def init_view(self):
        self.form = Opportunity_main_from
        # 顯示數據的頁面
        self.view_template_name = "PMIS/opportunity.html"

    @URL(url='Technical_Material')
    def technical(self, request):
        param = request.GET.get("param")
        inc_id = request.GET.get("inc_id")
        detailplanner = request.GET.get("detailplanner",'')
        tecmbQuerySet = None
        if inc_id!=None and  inc_id!='':
            tecmbQuerySet = list(models.VTecmb.objects.filter(inc_id=inc_id))
        else:
            tecmbQuerySet = list(models.VTecmb.objects.filter(mb023=param))
        technologyId = tecmbQuerySet[0].id
        ##記錄哪個人訪問這個文檔
        try:
            cache_name = "{0}:{1}".format(__name__, 'ViewTechnical')
            ViewInfo = cache.get(cache_name)
            if ViewInfo:
                ViewInfo.append({'contact':request.user.username, 'Date':datetime.datetime.now(), 'TechnicalId':tecmbQuerySet[0].mb023.strip()})
            else:
                ViewInfo = [{'contact':request.user.username, 'Date':datetime.datetime.now(), 'TechnicalId':tecmbQuerySet[0].mb023.strip()}]
            cache.set(cache_name, ViewInfo, timeout=60 * 60 * 24) #過期時間為1天    
        except Exception as e:
            pass
        fields = [field.attname for field in Tecmc._meta.get_fields()]
        fields.remove('mc006')
        tecmcQuerySet = list(Tecmc.objects.values(*fields).filter(id=technologyId))
        if tecmcQuerySet is None:
            return HttpResponse(status=404)
        steps_list = list()
        example_list = list()
        precaution_list = list()
        for i in tecmcQuerySet:
            if i['mc003'] == '1':
                steps_list.append(i)
            elif i['mc003'] == '2':
                example_list.append(i)
            elif i['mc003'] == '3':
                precaution_list.append(i)
        data = {'tecmb': tecmbQuerySet[0], 'steps': steps_list,
                'example': example_list, 'precaution': precaution_list}
        if detailplanner!='':
            return render(request, "PMIS/Technical_Material_show.html", data)
        return render(request, "PMIS/Technical_Material.html", data)


    @URL(url='Technical_Image')
    def getURL(self, request):
        id = request.GET.get("id")
        # 以-分割
        img = id.split('-')
        # 找到對應圖片數據
        image = list(Tecmc.objects.filter(
            id=img[0], mc003=img[1], mc004=img[2])[0:1])
        image = image[0].mc006
        if image is None:
            return HttpResponse(status=404)
        response = HttpResponse(image, content_type='image/jpg')
        response.__setitem__("Content-Disposition", "inline;filename="+id)

        return response

    @URL(url="tree")
    def ReportManagerjson(self, request):
        if request.method == 'GET':
            try:
                result = {}
                Reportfolders = list(Tecma.objects.all().values())
                result['state'] = 200
                result['data'] = Reportfolders
                return JsonResponse(result)
            except Exception as e:
                print(e)
                return JsonResponse(e) 


    @URL(url="getTechnicid")
    def getTechnicidjson(self, request):
        if request.method == 'GET':
            try:
                result = {}
                TechnicId = request.GET.get('TechnicId')
                Reportfolders = list(models.Tecmb.objects.filter(mb023__contains=TechnicId).values().order_by('-mb023')[:1])
                if len(Reportfolders)>0:
                    theTechnicid = "%05d"% (int(Reportfolders[0]['mb023'][8:13])+1)
                    maxTechnicid=[{'theTechnicid':TechnicId+theTechnicid}]
                else:
                    maxTechnicid=[{'theTechnicid':TechnicId+'00001'}]    
                result['state'] = True
                result['data'] = maxTechnicid
                return JsonResponse(result)
            except Exception as e:
                print(e)
                return JsonResponse(e) 

    @URL(url="updateTechnical")
    def updateTechnical(self, request):
        if request.method == 'POST':
            try:
                result = {'status':False,'msg':'','data':{}}
                thestatus = request.POST.get('status')
                technicstatus = 'N'
                if thestatus =='C':
                    technicstatus = 'T'
                if thestatus =='Y':
                    technicstatus = 'C'
                Technicalno = request.POST.get('Technicalno')
                if Technicalno:
                    Technicdata = models.Tecmb.objects.filter(mb023=Technicalno.strip(' '))
                    if Technicdata:
                        Technicdata.update(udf03=thestatus,mb020=technicstatus)   
                        result['status'] = True
                    else:
                        result['status'] =False
                        result['msg'] = '未找到對應技術文檔'
                return JsonResponse(result)
            except Exception as e:
                print(e)
                return JsonResponse(e)

    
    @URL(url='Get_technical')
    def Get_technical(self, request):
        result = {}
        param = request.GET.get("param")
        inc_id = request.GET.get("inc_id")
        get_img = request.GET.get("get_img",'')
        tecmbQuerySet = None
        if inc_id!='':
            tecmbQuerySet = list(models.VTecmb.objects.filter(inc_id=inc_id).values())
        else:
            tecmbQuerySet = list(models.VTecmb.objects.filter(mb023=param).values())
        technologyId = tecmbQuerySet[0]['id']
        mb015ctext = tecmbQuerySet[0]['mb015c']
        fields = [field.attname for field in Tecmc._meta.get_fields()]
        fields.remove('mc006')
        tecmcQuerySet = list(Tecmc.objects.values(*fields).filter(id=technologyId))
        if get_img!='':
            hasimg_list = Tecmc.objects.values('inc_id').filter(id=technologyId,mc006__isnull=True)
            hasimg_list = [item['inc_id'] for item in hasimg_list]
        if tecmcQuerySet is None:
            result['state'] = 404
            result['data'] = {}
            return JsonResponse(result)
        steps_list = list()
        example_list = list()
        precaution_list = list()
        for i in tecmcQuerySet:
            if get_img!='':
                i['hasimg'] = 'Y'
                if i['inc_id'] in hasimg_list:
                    i['hasimg'] = 'N'
            if i['mc003'] == '1':
                steps_list.append(i)
            elif i['mc003'] == '2':
                example_list.append(i)
            elif i['mc003'] == '3':
                precaution_list.append(i)
        data = {'steps': steps_list,'example': example_list, 'precaution': precaution_list,'mb015ctext':mb015ctext,'tecmb':tecmbQuerySet}
        result['state'] = 200
        result['data'] = data
        return JsonResponse(result)   

    @URL(url='Get_status')
    def Get_status(self, request):
        result = {}
        param = request.GET.get("param")
        tecmbQuerySet = list(models.VTecmb.objects.filter(mb023=param))
        technologyId = tecmbQuerySet[0].id
        fields = [field.attname for field in Tecmc._meta.get_fields()]
        fields.remove('mc006')
        tecmcQuerySet = list(Tecmc.objects.values(*fields).filter(id=technologyId))
        if tecmcQuerySet is None:
            result['state'] = 404
            result['data'] = {}
            return JsonResponse(result)
        steps_list = list()
        example_list = list()
        precaution_list = list()
        for i in tecmcQuerySet:
            if i['mc003'] == '1':
                steps_list.append(i)
            elif i['mc003'] == '2':
                example_list.append(i)
            elif i['mc003'] == '3':
                precaution_list.append(i)
        data = {'steps': steps_list,'example': example_list, 'precaution': precaution_list}
        result['state'] = 200
        result['data'] = data
        return JsonResponse(result)   




#技術文檔表格
class Technical_Datatable(DatatablesServerSideView):
    model = models.VTecmb
    columns =['mb023','mb004','mb005','mb008','parentid','mb001','mb006','mb015','mb016','inc_id','mb015c','mb019']
    searchable_columns = columns
    is_plus_search = True 

    def get_initial_queryset(self):
        ##記錄所有有問題的問題
        try:
            if self.request.GET.get("search[value]",""):
                cache_name = "{0}:{1}".format(__name__, 'TechnicalQuestion')
                Questions = cache.get(cache_name)
                if Questions:
                    Questions = list(filter(lambda x:x['contact'] != self.request.user.username or (x['contact'] == self.request.user.username and abs((datetime.datetime.now() - x['Date']).total_seconds()) > 10), Questions))
                    Questions.append({'contact':self.request.user.username, 'Date':datetime.datetime.now(), 'Question':self.request.GET.get("search[value]")})
                else:
                    Questions = [{'contact':self.request.user.username, 'Date':datetime.datetime.now(), 'Question':self.request.GET.get("search[value]")}]
                cache.set(cache_name, Questions, timeout=60 * 60 * 24) #過期時間為1天    
        except Exception as e:
            pass
        bdate = self.request.GET.get("bdate")
        edate = self.request.GET.get("edate")
        qs = self.model.objects.all()
        if bdate is not None and bdate != '' and edate is not None and edate != '':
            bdate = bdate.replace('-','')
            edate = edate.replace('-','')
            qs = qs.filter(create_date__range=(bdate, edate))      
        return qs

    def doit(self, request):
        # 检查是否传入 "ai_search" 参数
        search_text = request.GET.get("search[value]", "")
        if search_text and search_text.startswith("ai_search:"):
            http_method = request.GET
            if request.method == "POST":
                http_method = request.POST
            local = http_method.copy()
            local["search[value]"] = search_text.replace("ai_search:","")
            if request.method == 'POST':
                request.POST = local
            else:
                request.GET = local            
            return self.ai_search(request)
        else:
            return super().doit(request)

    def ai_search(self, request):
        """
        使用 search API 查询，然后根据返回的 id 从 vtecmb 表查询数据并合并 score_value
        """
        http_method = request.GET
        if request.method == "POST":
            http_method = request.POST        
        search_text = request.GET.get("search[value]", "")
        search_url = f"{settings.AI_RESTAPI_SERVER}/api/search/"

        # 构建请求数据
        payload = {
            "search_text": search_text,
            "search_type": "technical",
        }

        # 调用 search API
        try:
            response = AsyncioTools.async_fetch_http_json({'data':{"url":search_url, 'params':payload, 'method':"post"}})
            response_data = response['data']
            if response_data['most_similar_topic']:
                ai_result = [response_data['most_similar_topic']] + response_data['similar_topics']
            else:
                ai_result = []
            # 从 search API 响应中获取 id 列表
            ids = [int(topic['id']) for topic in ai_result]

            # 如果没有返回 id，直接返回空数据
            if not ids:
                return JsonResponse({"data": [], "recordsTotal": 0, "recordsFiltered": 0, "draw": int(http_method.get('draw', 0))})

            # 根据 ids 查询 vtecmb 表中的数据
            vtecmb_queryset = self.model.objects.filter(inc_id__in=ids)

            id_score_map = {int(item['id']): item['score_value'] for item in ai_result}
            # 合并 search API 中的 score_value
            for record in vtecmb_queryset:
                if record.inc_id in id_score_map:
                    setattr(record, "score_value", id_score_map[record.inc_id])
                else:
                    setattr(record, "score_value", 0)
            vtecmb_queryset = sorted(vtecmb_queryset, key=lambda x: x.score_value, reverse=True)
            # 将结果返回
            return self.build_json_response(vtecmb_queryset)

        except Exception as e:
            return JsonResponse({"error": str(e)}, status=500)

    def build_json_response(self, queryset):
        """
        构建符合 Datatables 的 JSON 响应
        """
        json_data = []

        for cur_object in queryset:
            retdict = {fieldname: self.render_column(cur_object, fieldname)
                       for fieldname in self.columns}
            retdict['DT_RowId'] = cur_object.pk
            retdict['score_value'] = getattr(cur_object, 'score_value', None)  # 添加 score_value 字段
            self.customize_row(retdict, cur_object)
            json_data.append(retdict)

        return JsonResponse({
            "data": json_data,
            "recordsTotal": len(queryset),
            "recordsFiltered": len(queryset),
            "draw": int(self.request.GET.get('draw', 0)),
        })   

    def prepare_results(self, qs):
        '''
        功能描述：將查詢返回的model數據解析出來
        '''
        json_data = []

        for cur_object in qs:
            retdict = {fieldname: self.render_column(cur_object, fieldname)
                       for fieldname in self.columns}
            retdict['DT_RowId'] = cur_object.pk
            retdict['score_value'] = 0
            self.customize_row(retdict, cur_object)
            json_data.append(retdict)
        return json_data             
    

def get_technical_paraphrase(request):
    # 获取前端传来的数据
    result = {'status':False, 'msg':'', "data":[]}
    id = request.GET.get("id")
    try:
        if not id:
            return JsonResponse(result, safe=False)
        else:
            url = f"{settings.AI_RESTAPI_SERVER}/api/search/get-data/"        
            params = {"search_type":"technical", "ids":[str(id)]}
            response = AsyncioTools.async_fetch_http_json({
            'data': {
                "url": url,
                'params': params,
                'method': "post"
            }})
            data = response['data']['data']
            if not data: #如果沒有找到數據，則查詢當前數據
                qs =  Tecmb.objects.values("inc_id", "mb004").filter(inc_id=id)[:1]
                data = [{"id":str(row['inc_id']), 'topic':row['mb004']} for row in qs]
            result['status'] = True
            result['data'] = data
    except Exception as e:
        print(str(e))
    return JsonResponse(result, safe=False)


def update_technical_embeddings(request, ids=None):
    """
    使用 update-technical-embeddings API 更新 embeddings
    """
    result = {'status':False, 'msg':"", 'data':{}}
    # 调用 API 更新 embeddings
    try:
        update_url = f"{settings.AI_RESTAPI_SERVER}/api/search/update-embeddings/" 
        if ids:
            technical_inc_ids = ids
            technical_inc_id = None
        else:
            request_data = json.loads(request.body)       
            technical_inc_id = request_data.get("id", "");
            technical_inc_ids = request_data.get("ids", [])

        if not technical_inc_id and not technical_inc_ids:
            return JsonResponse(result, safe=False)
        temp_inc_ids = [technical_inc_id] if technical_inc_id else technical_inc_ids
        qs =  Tecmb.objects.values("inc_id", "mb004").filter(inc_id__in=[str(inc_id) for inc_id in temp_inc_ids])        
        data = [{"id":str(row['inc_id']), 'topic':row['mb004']} for row in qs]
        if technical_inc_id:
            qs = Tecmb.objects.filter(inc_id=technical_inc_id)
            paraphrase_count = request_data.get('paraphrase_count', 3)
            gpt_model = request_data.get('gpt_model', 'gpt3')
            custom_paraphrases = request_data.get('custom_paraphrases', False)
            if custom_paraphrases:
                paraphrases_english = request_data.get('paraphrases_english', [])
                paraphrases_chinese = request_data.get('paraphrases_chinese', [])
                if not paraphrases_english and not paraphrases_chinese: #如果英文和中文都為空則設置custom_paraphrases為False
                    custom_paraphrases = False
                else:
                    if paraphrases_english:
                        data['topic_paraphrases_english'] = paraphrases_english
                    if paraphrases_chinese:
                        data['topic_paraphrases_chinese'] = paraphrases_chinese
            params = {'gpt_model':gpt_model, 'paraphrase_count':paraphrase_count, 'custom_paraphrases':custom_paraphrases, 'data':data}
        else:
            params = {'data':data}
        update_response = AsyncioTools.async_fetch_http_json({
            'data': {
                "url": update_url,
                'params': params,
                'method': "post"
            }
        })
        result = update_response['data']
    except Exception as e:
        print(str(e))
    return JsonResponse(result, safe=False)

#目錄表格
class Catalogue_Datatable(DatatablesServerSideView):
    model = models.VTecma
    columns =['ma001','ma002','ma003','ma003c','ma003b','ma007','inc_id']
    searchable_columns = columns


#目錄表新增類
class TecmaCreateView(SWCreateView):
    model = models.Tecma

    def get_initial(self, instance): 
        self.set_max_seqno(instance)  

    def set_max_seqno(self, instance):
        '''
        功能描述：獲取最大單號
        參數說明:
            instance:需要保存或初始化的model實例
        '''
        todaydate = datetime.datetime.now().strftime("%Y%m%d")
        rs = self.model.objects.filter(ma001__startswith = todaydate).aggregate(Max('ma001'))
        if rs and rs['ma001__max']:
            max_seq_no = str(int(rs['ma001__max']) + 1)
        else:
            max_seq_no = todaydate+'001'
        instance.ma001 = max_seq_no
    
    def set_ma007_value(self, instance):
        if (instance.ma007==None) and instance.ma002=='placeholder': 
            instance.ma007 = instance.ma001 
            instance.ma002 = ''
            instance.save()

    def save_other(self, instance):
        self.set_ma007_value(instance)  
         
    def save_check(self, instance):
        '''
        功能描述:保存前检查数据
        參數說明:
            instance:需要保存的model實例
        返回值:
            1) boolean 数据是否合法
            2) msg 数据不合法的原因
        '''
        if instance.ma002=='placeholder': 
            TempData = models.Tecma.objects.filter(ma002='',ma003=instance.ma003)
            if TempData.exists():
                return False,'已存在對應分類'
        if instance.ma002!='placeholder' and instance.ma002!='' and instance.ma002!='' \
        and instance.ma007!='' and instance.ma007!='' and instance.ma002==instance.ma007:     
            TempData = models.Tecma.objects.filter(ma002=instance.ma002,ma007=instance.ma007,ma003=instance.ma003)
            if TempData.exists():
                return False,'已存在對應區域'
        if instance.ma002!='placeholder' and instance.ma002!=None and instance.ma002!='' \
        and instance.ma007!=None and instance.ma007!='' and instance.ma002!=instance.ma007:  
            TempData = models.Tecma.objects.filter(ma002=instance.ma002,ma007=instance.ma007,ma003=instance.ma003)
            if TempData.exists():
                return False,'已存在對應子區域'
        return True, ''  


#目錄表修改類
class TecmaUpdateView(SWUpdateView):
    model = models.Tecma
    
    def set_ma007_value(self, instance):
        if instance.ma002=='placeholder': 
            instance.ma002 = ''
            instance.save()

    def save_other(self, instance):
        self.set_ma007_value(instance)   

    
    def save_check(self, instance):
        '''
        功能描述:保存前检查数据
        參數說明:
            instance:需要保存的model實例
        返回值:
            1) boolean 数据是否合法
            2) msg 数据不合法的原因
        '''
        TempData = models.Tecma.objects.filter(ma002=instance.ma002,ma007=instance.ma007,ma003=instance.ma003)
        if TempData.exists():
            return False,'已存在對應分類數據'
        return True, '' 


class TechnicalUpdateView(SWUpdateView):
    model = models.Tecmb

    def save_other(self, instance):
        instance = self.model.objects.get(inc_id = instance.inc_id)
        self.create_category(instance)
        Tecmb_UpdateTask(instance)
        #獲取下標值，將值轉為int類型並排序
        typefirstlist = set([re.match('typefirst\[(\d+)\]', key).group(1) for key in self.request.POST.keys() if re.match('typefirst\[\d+\][.]', key)])
        typesecondlist = set([re.match('typesecond\[(\d+)\]', key).group(1) for key in self.request.POST.keys() if re.match('typesecond\[\d+\][.]', key)])
        typethirdlist = set([re.match('typethird\[(\d+)\]', key).group(1) for key in self.request.POST.keys() if re.match('typethird\[\d+\][.]', key)])
        
        typefirstlist = [int(i) for i in typefirstlist]
        typefirstlist.sort()

        typesecondlist = [int(i) for i in typesecondlist]
        typesecondlist.sort()

        typethirdlist = [int(i) for i in typethirdlist]
        typethirdlist.sort()

        # print(typefirstlist)
        # print(typesecondlist)
        # print(typethirdlist)

        #獲取前端傳入的參數
        postData = self.request.POST
        for i in typefirstlist:
            strmc003 = "%04d"% (i)
            typefirstTable = models.Tecmc.objects.filter(id=instance.id,mc003='1',mc004=strmc003)
            if typefirstTable:
                updateSolutionData(strmc003,'1',postData,typefirstTable)
            else: 
                insertSolutionData(instance.id,postData,strmc003,'1')   
            typefirstTable = models.Tecmc.objects.filter(id=instance.id,mc003='1',mc004=strmc003)
            if (postData.get('typefirst['+strmc003+'].deleteimg')=='') or (postData.get('typefirst['+strmc003+'].deleteimg')=='on'):
                typefirstTable.update(mc006=None)               
            if (postData.get('typefirst['+strmc003+'].deletedata')=='') or (postData.get('typefirst['+strmc003+'].deletedata')=='on'):
                typefirstTable.delete()

        for i in typesecondlist:
            strmc003 = "%04d"% (i)
            typesecondlist = models.Tecmc.objects.filter(id=instance.id,mc003='2',mc004=strmc003)
            if typesecondlist:
                updateSolutionData(strmc003,'2',postData,typesecondlist)
            else: 
                insertSolutionData(instance.id,postData,strmc003,'2')   
            typesecondlist = models.Tecmc.objects.filter(id=instance.id,mc003='2',mc004=strmc003)
            if (postData.get('typesecond['+strmc003+'].deletedata')=='') or (postData.get('typesecond['+strmc003+'].satisfactory')=='on'):
                typesecondlist.delete()
            if (postData.get('typesecond['+strmc003+'].deleteimg')=='') or (postData.get('typesecond['+strmc003+'].deleteimg')=='on'):
                typesecondlist.update(mc006=None)   

        for i in typethirdlist:
            strmc003 = "%04d"% (i)
            typethirdlist = models.Tecmc.objects.filter(id=instance.id,mc003='3',mc004=strmc003)
            if typethirdlist:
                updateSolutionData(strmc003,'3',postData,typethirdlist)
            else: 
                insertSolutionData(instance.id,postData,strmc003,'3')   
            typethirdlist = models.Tecmc.objects.filter(id=instance.id,mc003='3',mc004=strmc003)
            if (postData.get('typethird['+strmc003+'].deletedata')=='') or (postData.get('typethird['+strmc003+'].satisfactory')=='on'):
                typethirdlist.delete()
            if (postData.get('typethird['+strmc003+'].deleteimg')=='') or (postData.get('typethird['+strmc003+'].deleteimg')=='on'):
                typethirdlist.update(mc006=None)   
        Temp_tecmb = models.VTecmb.objects.get(inc_id = instance.inc_id)
        instance.mb015c = Temp_tecmb.mb015c
        service = ElasticService()
        service.UpdateElastic(instance)
        
    
    def save_check(self, instance):
        '''
        功能描述:保存前检查数据
        參數說明:
            instance:需要保存的model實例
        返回值:
            1) boolean 数据是否合法
            2) msg 数据不合法的原因
        '''
        # print(self.request.POST.get('mb015c'))
        # VTecmadata = list(VTecma.objects.filter(ma003=self.request.POST.get('mb015c')).values()[:1])
        # if len(VTecmadata)>0:
        #     instance.mb015 = VTecmadata[0]['ma001']
        check_status, check_msg = True,''
        # if (instance.mb008 == None) or (instance.mb008.strip( ' ' )) == '':
        #     check_status, check_msg = False ,'作用不能為空'
        # if (instance.mb001 == None) or (instance.mb001.strip( ' ' )) == '':
        #     check_status, check_msg = False ,'請選擇分類'
        # if (instance.mb016 == None) or (len(instance.mb016)< 3):
        #     check_status, check_msg = False ,'Area長度必須大於等於3'
        # if (instance.mb016 == None) or (instance.mb016.strip( ' ' )) == '':
        #     check_status, check_msg = False ,'Area不能為空'        
        if (instance.mb005 == None) or (instance.mb005.strip( ' ' )) == '':
            check_status, check_msg = False ,'聯繫人不能為空'
        if (instance.mb004 == None) or (instance.mb004.strip( ' ' )) == '':
            check_status, check_msg = False ,'主題不能為空'
        return check_status, check_msg       



    def create_category(self,instance):
        strmb015c = self.request.POST.get('mb015c')
        TempData_category = models.Tecma.objects.filter(ma001=instance.mb015,ma003=strmb015c)
        if not TempData_category.exists():
            Table_Tecma = models.Tecma()
            ModelTools.set_basic_field_info(self.request, models.Tecma, Table_Tecma)
            self.set_max_tecma_seqno(Table_Tecma)
            Table_Tecma.ma002=''
            Table_Tecma.ma003=strmb015c
            Table_Tecma.ma007=Table_Tecma.ma001
            Table_Tecma.save()
            instance.mb015 = Table_Tecma.ma001
            instance.save()
        if  (instance.mb015!=None and instance.mb015!='') and (instance.mb016!=None and instance.mb016!=''):
            TempData_area = models.Tecma.objects.filter(ma002=instance.mb015,ma003=instance.mb016)
            if not TempData_area.exists():
                Table_Tecma = models.Tecma()
                ModelTools.set_basic_field_info(self.request, models.Tecma, Table_Tecma)
                self.set_max_tecma_seqno(Table_Tecma)
                Table_Tecma.ma002=instance.mb015
                Table_Tecma.ma003=instance.mb016
                Table_Tecma.ma007=instance.mb015
                Table_Tecma.save()

            if (instance.mb026!=None and instance.mb026!=''):
                TempData_area = models.Tecma.objects.filter(ma002=instance.mb015,ma003=instance.mb016)
                if TempData_area.exists():
                    TempData_area = TempData_area.values()[0]
                    TempData_subarea = models.Tecma.objects.filter(ma002=TempData_area['ma001'],ma003=instance.mb026,ma007=instance.mb015)
                    if not TempData_subarea.exists():
                        Table_Tecma = models.Tecma()
                        ModelTools.set_basic_field_info(self.request, models.Tecma, Table_Tecma)
                        self.set_max_tecma_seqno(Table_Tecma)
                        Table_Tecma.ma002=TempData_area['ma001']
                        Table_Tecma.ma003=instance.mb026
                        Table_Tecma.ma007=instance.mb015
                        Table_Tecma.save()

    
    def set_max_tecma_seqno(self, instance):
        '''
        功能描述：獲取最大單號
        參數說明:
            instance:需要保存或初始化的model實例
        '''
        todaydate = datetime.datetime.now().strftime("%Y%m%d")
        rs = models.Tecma.objects.filter(ma001__startswith = todaydate).aggregate(Max('ma001'))
        if rs and rs['ma001__max']:
            max_seq_no = str(int(rs['ma001__max']) + 1)
        else:
            max_seq_no = todaydate+'001'
        instance.ma001 = max_seq_no



class TechnicalCreateView(SWCreateView):
    model = models.Tecmb

    def get_initial(self, instance): 
        instance.mb006 = datetime.datetime.now().strftime("%Y%m%d")
        instance.mb020 = 'N'
        instance.create_date = datetime.datetime.now().strftime("%Y%m%d")
        self.set_max_seqno(instance)  
        self.Set_DefaultData(instance)

    def set_max_seqno(self, instance):
        '''
        功能描述：獲取最大單號
        參數說明:
            instance:需要保存或初始化的model實例
        '''
        rs = self.model.objects.all().aggregate(Max('id'))
        if len(rs) > 0:
            max_seq_no = rs['id__max'] + 1
        else:
            max_seq_no = 1
        instance.id = max_seq_no

    def save_other(self, instance):
        instance = self.model.objects.get(inc_id = instance.inc_id)
        self.create_category(instance)
        Tecmb_UpdateTask(instance)
        #獲取下標值，將值轉為int類型並排序
        typefirstlist = set([re.match('typefirst\[(\d+)\]', key).group(1) for key in self.request.POST.keys() if re.match('typefirst\[\d+\][.]', key)])
        typesecondlist = set([re.match('typesecond\[(\d+)\]', key).group(1) for key in self.request.POST.keys() if re.match('typesecond\[\d+\][.]', key)])
        typethirdlist = set([re.match('typethird\[(\d+)\]', key).group(1) for key in self.request.POST.keys() if re.match('typethird\[\d+\][.]', key)])
        
        typefirstlist = [int(i) for i in typefirstlist]
        typefirstlist.sort()

        typesecondlist = [int(i) for i in typesecondlist]
        typesecondlist.sort()

        typethirdlist = [int(i) for i in typethirdlist]
        typethirdlist.sort()

        #獲取前端傳入的參數
        postData = self.request.POST
        for i in typefirstlist:
            strmc003 = "%04d"% (i)
            insertSolutionData(instance.id,postData,strmc003,'1')   
            typefirstTable = models.Tecmc.objects.filter(id=instance.id,mc003='1',mc004=strmc003)    
            if (postData.get('typefirst['+strmc003+'].deleteimg')=='') or (postData.get('typefirst['+strmc003+'].deleteimg')=='on'):
                typefirstTable.update(mc006=None)               
            if (postData.get('typefirst['+strmc003+'].deletedata')=='') or (postData.get('typefirst['+strmc003+'].deletedata')=='on'):
                typefirstTable.delete()

        for i in typesecondlist:
            strmc003 = "%04d"% (i)
            insertSolutionData(instance.id,postData,strmc003,'2')   
            typesecondlist = models.Tecmc.objects.filter(id=instance.id,mc003='2',mc004=strmc003)
            if (postData.get('typesecond['+strmc003+'].deletedata')=='') or (postData.get('typesecond['+strmc003+'].satisfactory')=='on'):
                typesecondlist.delete()
            if (postData.get('typesecond['+strmc003+'].deleteimg')=='') or (postData.get('typesecond['+strmc003+'].deleteimg')=='on'):
                typesecondlist.update(mc006=None)   

        for i in typethirdlist:
            strmc003 = "%04d"% (i)
            insertSolutionData(instance.id,postData,strmc003,'3')   
            typethirdlist = models.Tecmc.objects.filter(id=instance.id,mc003='3',mc004=strmc003)
            if (postData.get('typethird['+strmc003+'].deletedata')=='') or (postData.get('typethird['+strmc003+'].satisfactory')=='on'):
                typethirdlist.delete()
            if (postData.get('typethird['+strmc003+'].deleteimg')=='') or (postData.get('typethird['+strmc003+'].deleteimg')=='on'):
                typethirdlist.update(mc006=None)   
        Temp_tecmb = models.VTecmb.objects.get(inc_id = instance.inc_id)
        instance.mb015c = Temp_tecmb.mb015c
        service = ElasticService()
        service.UpdateElastic(instance)                
        new_mb004 = instance.mb004.strip() if instance.mb004 else ""
        if new_mb004:
            thread = threading.Thread(target=update_technical_embeddings, args=(None, [instance.inc_id]))
            thread.start()         
                
    def save_check(self, instance):
        '''
        功能描述:保存前检查数据
        參數說明:
            instance:需要保存的model實例
        返回值:
            1) boolean 数据是否合法
            2) msg 数据不合法的原因
        '''
        # print(self.request.POST.get('mb015c'))
        # VTecmadata = list(VTecma.objects.filter(ma003=self.request.POST.get('mb015c')).values()[:1])
        # if len(VTecmadata)>0:
        #     instance.mb015 = VTecmadata[0]['ma001']
        check_status, check_msg = Tecmb_save_check(instance)        
        if check_status:
            # self.Set_DefaultData(instance)    
            self.set_mb023(instance)
        return check_status, check_msg
    
    def set_mb023(self, instance):
        if (instance.mb015!='' or instance.mb015!=None) and (instance.mb016!='' or instance.mb016!=None) and (instance.mb023!='' or instance.mb023!=None):
            classify = models.Tecma.objects.filter(ma001=instance.mb015)
            if classify.exists():
                classify = classify.values('ma003')[0]['ma003'][0:3]
                TechnicId = f'{classify}-{instance.mb016[0:3]}-'
                Reportfolders = list(Tecmb.objects.filter(mb023__contains=TechnicId).values().order_by('-mb023')[:1])
                if len(Reportfolders)>0:
                    theTechnicid = "%05d"% (int(Reportfolders[0]['mb023'][8:13])+1)
                    instance.mb023=TechnicId+theTechnicid
                else:
                    instance.mb023=TechnicId+'00001'   


    #設置默認的分類和area
    def Set_DefaultData(self, instance):
        default_classify = models.Syspara.objects.filter(nfield='tecDefaultCategory',ftype='tecDefaultCategory')
        if default_classify.exists():
            ma001 = default_classify.values('fvalue')[0]['fvalue'].replace(' ','')
            default_classify = models.Tecma.objects.filter(ma001=ma001)
            default_area = models.Tecma.objects.filter(ma002=ma001,ma003__icontains='area')
            if default_classify.exists() and (instance.mb015=='' or instance.mb015==None):
                default_classify = default_classify.values('ma001','ma007','ma003')[0]
                instance.mb001 = default_classify['ma003']  
                instance.mb015 = default_classify['ma001']
            if default_area.exists() and (instance.mb016=='' or instance.mb016==None):
                instance.mb016 = default_area.values('ma001','ma003')[0]['ma003']


    def create_category(self,instance):
        strmb015c = self.request.POST.get('mb015c')
        TempData_category = models.Tecma.objects.filter(ma002__exact='',ma003=strmb015c)
        if not TempData_category.exists():
            Table_Tecma = models.Tecma()
            ModelTools.set_basic_field_info(self.request, models.Tecma, Table_Tecma)
            self.set_max_tecma_seqno(Table_Tecma)
            Table_Tecma.ma002=''
            Table_Tecma.ma003=strmb015c
            Table_Tecma.ma007=Table_Tecma.ma001
            Table_Tecma.save()
            instance.mb015 = Table_Tecma.ma001
            instance.save()

        if TempData_category.exists() and (instance.mb015==None or instance.mb015==''):
            instance.mb015 = TempData_category['ma001']
            instance.save()

        if  (instance.mb015!=None and instance.mb015!='') and (instance.mb016!=None and instance.mb016!=''):
            TempData_area = models.Tecma.objects.filter(ma002=instance.mb015,ma003=instance.mb016)
            if not TempData_area.exists():
                Table_Tecma = models.Tecma()
                ModelTools.set_basic_field_info(self.request, models.Tecma, Table_Tecma)
                self.set_max_tecma_seqno(Table_Tecma)
                Table_Tecma.ma002=instance.mb015
                Table_Tecma.ma003=instance.mb016
                Table_Tecma.ma007=instance.mb015
                Table_Tecma.save()

            if (instance.mb026!=None and instance.mb026!=''):
                TempData_area = models.Tecma.objects.filter(ma002=instance.mb015,ma003=instance.mb016)
                if TempData_area.exists():
                    TempData_area = TempData_area.values()[0]
                    TempData_subarea = models.Tecma.objects.filter(ma002=TempData_area['ma001'],ma003=instance.mb026,ma007=instance.mb015)
                    if not TempData_subarea.exists():
                        Table_Tecma = models.Tecma()
                        ModelTools.set_basic_field_info(self.request, models.Tecma, Table_Tecma)
                        self.set_max_tecma_seqno(Table_Tecma)
                        Table_Tecma.ma002=TempData_area['ma001']
                        Table_Tecma.ma003=instance.mb026
                        Table_Tecma.ma007=instance.mb015
                        Table_Tecma.save()

    
    def set_max_tecma_seqno(self, instance):
        '''
        功能描述：獲取最大單號
        參數說明:
            instance:需要保存或初始化的model實例
        '''
        todaydate = datetime.datetime.now().strftime("%Y%m%d")
        rs = models.Tecma.objects.filter(ma001__startswith = todaydate).aggregate(Max('ma001'))
        if rs and rs['ma001__max']:
            max_seq_no = str(int(rs['ma001__max']) + 1)
        else:
            max_seq_no = todaydate+'001'
        instance.ma001 = max_seq_no

#更新Tecmc表數據
def updateSolutionData(strmc003,thetype:str,postData,Tabledata:object):
    '''
    功能描述：更新Tecmc表數據
    '''
    #獲取字段值
    typeindex = ''
    if thetype == '1':
        typeindex = 'typefirst'
    if thetype == '2':
        typeindex = 'typesecond'
    if thetype == '3':  
        typeindex = 'typethird'
    mc006 = postData.get(typeindex+'['+strmc003+'].imginput','')
    if '/PMIS/opportunity/Technical_Image' in mc006:
        mc006=''
    if mc006!=None and mc006!='':
        mc006=base64.b64decode(mc006)         
    strmc005 = postData.get(typeindex+'['+strmc003+'].mc005','')
    if mc006!='':
        Tabledata.update(mc005=strmc005,mc006=mc006,mc007=strmc003)   
    else:    
        Tabledata.update(mc005=strmc005,mc007=strmc003)    

#新增Tecmc表數據
def insertSolutionData(incid:int,postData:object,strmc003,thetype:str):
    '''
    功能描述：新增Tecmc表數據
    '''
    typeindex = ''
    if thetype == '1':
        typeindex = 'typefirst'
    if thetype == '2':
        typeindex = 'typesecond'
    if thetype == '3':  
        typeindex = 'typethird'
    mc001 = datetime.datetime.now().strftime("%Y%m%d")
    mc002 = ''
    mc004 = set_max_solution(incid,thetype)
    mc005 = postData.get(typeindex+'['+strmc003+'].mc005','')
    mc006 = postData.get(typeindex+'['+strmc003+'].imginput','')
    if '/PMIS/opportunity/Technical_Image' in mc006:
        mc006=''
    if mc006!=None and mc006!='':
        mc006=base64.b64decode(mc006)     
    create_date = datetime.datetime.now().strftime("%Y%m%d%H%M%S")
    if mc006!='':
        solutionData = models.Tecmc(mc001=mc001,mc002=mc002,mc003=thetype,mc004=mc004,mc005=mc005,create_date=create_date,id=incid,mc006=mc006,mc007=strmc003)
    else:    
        solutionData = models.Tecmc(mc001=mc001,mc002=mc002,mc003=thetype,mc004=mc004,mc005=mc005,create_date=create_date,id=incid,mc007=strmc003)
    solutionData.save()

#獲取序號最大值
def set_max_solution(incid,thetype):
    #獲取序號最大值
    rs = models.Tecmc.objects.values('mc004').filter(id=incid,mc003=thetype).order_by('-mc004')[:1]
    if len(rs) > 0:
        max_seq_no = "%04d"% (int(rs[0]['mc004'])+1)
    else:
        max_seq_no = '0001'
    return max_seq_no    


def Tecmb_save_check(instance):
    '''
    功能描述:保存前检查数据
    參數說明:
        instance:需要保存的model實例
    返回值:
        1) boolean 数据是否合法
        2) msg 数据不合法的原因
    '''
    check_status, check_msg = True,''
    # if (instance.mb008 == None) or (instance.mb008.strip( ' ' )) == '':
    #     check_status, check_msg = False ,'作用不能為空'
    # if (instance.mb001 == None) or (instance.mb001.strip( ' ' )) == '':
    #     check_status, check_msg = False ,'請選擇分類'
    # if (instance.mb016 == None) or (len(instance.mb016)< 3):
    #     check_status, check_msg = False ,'Area長度必須大於等於3'
    # if (instance.mb016 == None) or (instance.mb016.strip( ' ' )) == '':
    #     check_status, check_msg = False ,'Area不能為空'
    if (instance.mb005 == None) or (instance.mb005.strip( ' ' )) == '':
        check_status, check_msg = False ,'聯繫人不能為空'
    if (instance.mb004 == None) or (instance.mb004.strip( ' ' )) == '':
        check_status, check_msg = False ,'主題不能為空'
    return check_status, check_msg


def get_categorytree(request):
    '''
    功能描述：獲取技術文檔分類樹狀圖
    '''
    result = {'status':False, 'msg':'', 'data':[]}
    try:
        category = request.GET.get('category','')
        area = request.GET.get('area','')
        subarea = request.GET.get('subarea','')
        if subarea=='' and area=='' and category=='':
            TempData = models.VTecma.objects.all()  
            if TempData.exists():
                result['data'] = list(TempData.values())
        else:
            Temp_Data = []
            childrenNode=[]
            category_list=[]
            area_list=[]

            if category!='':
                Temp_Data = models.VTecma.objects.filter(Q(ma003__contains=category) & Q(ma002__exact=''))
                Temp_Data=list(Temp_Data.values())
                category_list = [item['ma001'].replace(' ','') for item in Temp_Data]
                childrenNode = list(models.VTecma.objects.filter(Q(ma007__in=category_list)).values())
            if area!='':
                Temp_Data = models.VTecma.objects.filter(Q(ma003__contains=area) & Q(ma002=F('ma007')) & ~Q(ma002__exact=''))
                if len(category_list)>0:
                    Temp_Data = Temp_Data.filter(ma002__in=category_list)
                Temp_Data=list(Temp_Data.values())
                area_list = [item['ma001'].replace(' ','') for item in Temp_Data]
                childrenNode = list(models.VTecma.objects.filter(Q(ma002__in=area_list)).values())
            if subarea!='':
                Temp_Data = models.VTecma.objects.filter(Q(ma003__contains=subarea) & ~Q(ma002=F('ma007')) & ~Q(ma002__exact=''))
                if len(category_list)>0:
                    Temp_Data = Temp_Data.filter(ma007__in=category_list)
                if len(area_list)>0:
                    Temp_Data = Temp_Data.filter(ma002__in=area_list)
                Temp_Data=list(Temp_Data.values())

            category_list = [item['ma007'].replace(' ','') for item in Temp_Data if item['ma007'].replace(' ','')!=None and item['ma007'].replace(' ','')!='']
            area_list = [item['ma002'].replace(' ','') for item in Temp_Data if item['ma002'].replace(' ','')!=None and item['ma002'].replace(' ','')!='']  
            TempIDList =  list(set(category_list + area_list))
            if len(childrenNode)>0:
                Temp_Data = Temp_Data+childrenNode
                Temp_Data = distinct(Temp_Data,['ma001'])    
            parentData = models.VTecma.objects.filter(ma001__in = TempIDList)  
            if parentData.exists():
                Temp_Data =  Temp_Data + list(parentData.values())
                Temp_Data = distinct(Temp_Data,['ma001'])    
            result['data'] = Temp_Data
        result['status'] = True
    except Exception as e:
        print(str(e))
        result['status'] = False
    return JsonResponse(result, safe=False)   

#目錄表新增類
class TecmfCreateView(SWCreateView):
    model = models.Tecmf

    def set_mf007_value(self, instance):
        if (instance.mf007==None) and instance.mf002=='#': 
            instance.mf007 = instance.inc_id 
            instance.save()

    def save_other(self, instance):
        self.set_mf007_value(instance)    
       

#目錄表修改類
class TecmfUpdateView(SWUpdateView):
    model = models.Tecmf


#保存技術文檔時更新或新增任務
def Tecmb_UpdateTask(instance):
    Pid = instance.pid
    Tid = instance.tid
    UserName = instance.mb005
    if UserName!='':
        UserName=UserName.replace(' ','')
    TaskId = instance.taskid
    Temp_task = models.Task.objects.filter(tid=Tid,pid=Pid,taskid=TaskId)
    if not Temp_task.exists():
        Temp_pumsut = models.Pmsut.objects.filter(ut001=UserName)
        if not Temp_pumsut.exists():
            return False,'當前用戶沒有設置session，請設置'  
        Temp_pumsut = Temp_pumsut.values()[0]
        Pid = Temp_pumsut['ut002'] 
        Tid = Temp_pumsut['ut003']
        TaskId = 10
        TempData = models.Task.objects.filter(pid=Pid,tid=Tid).aggregate(Max('taskid'))
        if TempData and TempData['taskid__max']:
            TaskId = TaskId+TempData['taskid__max']
        Temp_task = models.Task()
        Temp_task.pid=Pid
        Temp_task.tid=Tid
        Temp_task.taskid=TaskId
        Temp_task.planbdate=datetime.datetime.now()
        Temp_task.planedate=datetime.datetime.now()
        Temp_task.bdate=datetime.datetime.now()
        Temp_task.edate=datetime.datetime.now()
        Temp_task.progress='C'
        Temp_task.tasktype=170
        Temp_task.subtasktype=1740
        Temp_task.diff=1
        Temp_task.task=instance.mb004
        Temp_task.contact=UserName.strip()
        Temp_task.save()
        instance.pid = Temp_task.pid
        instance.tid = Temp_task.tid
        instance.taskid = Temp_task.taskid
        instance.save()
    else:
        Temp_task.update(task=instance.mb004,contact=instance.mb005)

    return True,''



def distinct(original_list, fields: list):  # 列表去重
    '''
        next(iter, default) iter是可迭代對象, 當迭代後沒有返回值, 則返回default定義的值
        all(iter) 所有元素不为0、''、False或者iterable为空,all(iterable)返回True,否则返回False
    '''
    if type(fields) == str:  # 當傳入的為單個字段時
        fields = [fields]
    unique = []
    for item in original_list:
        existing_item = next((d for d in unique if all(
            item[field] == d[field] for field in fields)), None)
        if existing_item is None:  # 不存在重複則加入
            unique.append(item)
    return unique




